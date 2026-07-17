"""
文档加载器 - 支持 PDF / Word / Markdown / TXT
"""
from pathlib import Path
from typing import List
from loguru import logger
import pymupdf  # fitz


def _table_to_markdown(pymupdf_table) -> str:
    """
    将 PyMuPDF Table 对象转为干净的 Markdown 表格。

    优先级：table.extract() → to_markdown() 兜底。
    extract() 对合并单元格的处理更好（重复值而非乱码 ColN）。
    """
    # ── 优先用 extract() 手动构建，避免 to_markdown() 的合并单元格乱码 ──
    try:
        data = pymupdf_table.extract()
        if data:
            lines = []
            for i, row in enumerate(data):
                cells = []
                for c in row:
                    cell_text = str(c).replace("\n", " ").replace("|", "/").strip() if c else ""
                    # 过滤 PyMuPDF 合并单元格生成的 ColN 占位符
                    if cell_text and not cell_text.startswith("Col"):
                        cells.append(cell_text)
                    else:
                        cells.append("")
                # 跳过全是 ColN 占位符的行
                if any(c for c in cells if c):
                    lines.append("| " + " | ".join(cells) + " |")
                    if len(lines) == 1:  # 第一个有效行后加分隔行
                        lines.append("|" + "|".join(["---"] * len(cells)) + "|")
            if len(lines) >= 3:  # 至少有表头+分隔行+一行数据
                return "\n".join(lines)
    except Exception:
        pass

    # ── 兜底：to_markdown() ──
    try:
        md = pymupdf_table.to_markdown()
        md = md.replace("<br>", " ")
        # 清理 ColN 占位符（合并单元格残留）
        import re
        md = re.sub(r'\bCol\d+\b', '', md)
        return md
    except Exception:
        return ""


def load_pdf(file_path: str) -> List[dict]:
    """
    加载 PDF 文件，按页提取文本 + 结构化表格

    V8.0 重构：用 "dict" 模式 + 坐标过滤，替代原来的 "blocks" 模式。
    - 页眉页脚（y < 8% 或 y > 92%）→ 过滤
    - 表格区域（find_tables 检测到的） → 独立存为 Markdown，文本中不重复
    - 相邻文本块按 y 坐标合并为段落

    返回: [{
        "text": "段落文本...",
        "tables": [{"markdown": "...", "rows": 5, "cols": 3}, ...],
        "page": 1,
        "source": "xxx.pdf"
    }, ...]
    """
    docs = []
    total_tables = 0
    try:
        doc = pymupdf.open(file_path)
        for page_num, page in enumerate(doc, start=1):
            # ── 1. 获取页面尺寸 ──
            page_height = page.rect.height
            header_y = page_height * 0.08   # 页面顶部 8% 视为页眉区
            footer_y = page_height * 0.92   # 页面底部 8% 视为页脚区

            # ── 2. 检测表格区域（用于后续文本过滤）──
            table_bboxes = []
            tables = []
            try:
                found = page.find_tables()
                for t in found.tables:
                    md = _table_to_markdown(t)
                    if md.strip():
                        tables.append({
                            "markdown": md,
                            "rows": t.row_count,
                            "cols": t.col_count,
                        })
                        table_bboxes.append(t.bbox)  # (x0, y0, x1, y1)
                total_tables += len(tables)
            except Exception as e:
                logger.debug(f"表格检测失败 P{page_num}: {e}")

            # ── 3. 用 "dict" 模式提取结构化文本 ──
            text_dict = page.get_text("dict")
            text_blocks = []

            for block in text_dict.get("blocks", []):
                # 只处理文本块（type=0），跳过图片块（type=1）
                if block.get("type") != 0:
                    continue

                bbox = block.get("bbox", (0, 0, 0, 0))
                y0, y1 = bbox[1], bbox[3]

                # 规则1: 过滤页眉页脚区域
                if y1 < header_y or y0 > footer_y:
                    continue

                # 规则2: 过滤表格区域内的文本（表格已单独提取）
                in_table = False
                for tb in table_bboxes:
                    if (y0 >= tb[1] - 5 and y1 <= tb[3] + 5 and
                        bbox[0] >= tb[0] - 5 and bbox[2] <= tb[2] + 5):
                        in_table = True
                        break
                if in_table:
                    continue

                # 提取块内文本
                lines = []
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                    line_text = line_text.strip()
                    if line_text:
                        lines.append(line_text)

                if lines:
                    text_blocks.append({
                        "y0": y0,
                        "text": "".join(lines),
                    })

            # ── 4. 规则3: 过滤纯页码行（短纯数字，在页眉页脚区域附近）──
            filtered_blocks = []
            for block in text_blocks:
                text = block["text"]
                # 纯数字且 ≤3 字符 → 疑似页码
                if text.isdigit() and len(text) <= 3:
                    continue
                # 纯数字+空格（如"4 5"） → 疑似页码
                if all(c.isdigit() or c.isspace() for c in text) and len(text.replace(" ", "")) <= 3:
                    continue
                filtered_blocks.append(block)

            # ── 5. 合并相邻文本块为段落（y 坐标接近的连续块）──
            paragraphs = []
            if filtered_blocks:
                current = filtered_blocks[0]["text"]
                for i in range(1, len(filtered_blocks)):
                    y_gap = filtered_blocks[i]["y0"] - filtered_blocks[i-1]["y0"]
                    # y 间距 < 20 点（约 0.7mm）视为同一段落
                    if y_gap < 20:
                        current += filtered_blocks[i]["text"]
                    else:
                        paragraphs.append(current)
                        current = filtered_blocks[i]["text"]
                paragraphs.append(current)

            text = "\n\n".join(paragraphs)

            # ── 6. 保留非空页面 ──
            if text.strip() or tables:
                docs.append({
                    "text": text.strip(),
                    "tables": tables,
                    "page": page_num,
                    "source": Path(file_path).name,
                })

        logger.info(
            f"PDF 加载完成: {file_path} -> {len(docs)} 页, {total_tables} 张表格"
        )
    except Exception as e:
        logger.error(f"PDF 加载失败: {e}")
        raise
    return docs


def load_docx(file_path: str) -> List[dict]:
    """
    加载 Word 文档，按段落分组为逻辑页。

    修复原因：之前把所有段落拼成一个超长文本塞到 page:1，
    导致 semantic_splitter 无法有效切分，整份研报只产出 1 个 chunk。
    现在按 ~15 段一组拆分成多个逻辑页，每页文本量适中，切分器能正常工作。
    """
    from docx import Document
    docs = []
    try:
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            logger.warning(f"DOCX 无有效文本: {file_path}")
            return docs

        # 按段落分组：每 15 段为一"页"，避免单页文本过短或过长
        PARAS_PER_PAGE = 15
        for i in range(0, len(paragraphs), PARAS_PER_PAGE):
            chunk_paras = paragraphs[i:i + PARAS_PER_PAGE]
            docs.append({
                "text": "\n".join(chunk_paras),
                "tables": [],
                "page": i // PARAS_PER_PAGE + 1,
                "source": Path(file_path).name,
            })

        logger.info(f"DOCX 加载完成: {file_path} → {len(paragraphs)} 段, {len(docs)} 逻辑页")
    except Exception as e:
        logger.error(f"DOCX 加载失败: {e}")
        raise
    return docs


def load_markdown(file_path: str) -> List[dict]:
    """加载 Markdown 文件"""
    docs = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        if text.strip():
            docs.append({
                "text": text.strip(),
                "tables": [],
                "page": 1,
                "source": Path(file_path).name,
            })
        logger.info(f"Markdown 加载完成: {file_path}")
    except Exception as e:
        logger.error(f"Markdown 加载失败: {e}")
        raise
    return docs


def load_txt(file_path: str) -> List[dict]:
    """加载纯文本文件"""
    docs = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        if text.strip():
            docs.append({
                "text": text.strip(),
                "tables": [],
                "page": 1,
                "source": Path(file_path).name,
            })
    except UnicodeDecodeError:
        # 尝试 GBK 编码（常见于中文文档）
        with open(file_path, "r", encoding="gbk") as f:
            text = f.read()
        if text.strip():
            docs.append({
                "text": text.strip(),
                "tables": [],
                "page": 1,
                "source": Path(file_path).name,
            })
    return docs


def load_document(file_path: str) -> List[dict]:
    """
    根据文件类型自动选择加载器
    """
    ext = Path(file_path).suffix.lower()
    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".doc": load_docx,
        ".md": load_markdown,
        ".txt": load_txt,
    }
    loader = loaders.get(ext)
    if not loader:
        raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {list(loaders.keys())}")
    return loader(file_path)
